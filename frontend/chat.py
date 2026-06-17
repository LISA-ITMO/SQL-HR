import streamlit as st
import requests
import os
import time
import uuid
import base64
import io
import re
import zipfile
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SERVER_URL = os.getenv("SERVER_URL")
POLL_INTERVAL = float(os.getenv("POLL_INTERVAL", "1.5"))
EXECUTOR = ThreadPoolExecutor(max_workers=1)
SYSTEM_NAME = 'Агент "Рекрутер"'
SYSTEM_SUBTITLE = "Подбор персонала в исполнительные органы государственной власти Санкт-Петербурга"
LOGO_PATH = Path(__file__).resolve().parent / "images" / "logo.png"
PROFILE_TEMPLATE_PATH = Path(__file__).resolve().parent / "profile.docx"
LOGO_DATA_URI = ""

if LOGO_PATH.exists():
    encoded_logo = base64.b64encode(LOGO_PATH.read_bytes()).decode("ascii")
    LOGO_DATA_URI = f"data:image/png;base64,{encoded_logo}"

if "messages" not in st.session_state:
    st.session_state.messages = []

if "session_id" not in st.session_state:
    st.session_state.session_id = None
if "pending_future" not in st.session_state:
    st.session_state.pending_future = None
if "stop_requested" not in st.session_state:
    st.session_state.stop_requested = False
if "last_polled_candidates" not in st.session_state:
    st.session_state.last_polled_candidates = []
if "final_fragment_message" not in st.session_state:
    st.session_state.final_fragment_message = None

st.set_page_config(
    page_title=SYSTEM_NAME,
    page_icon=str(LOGO_PATH) if LOGO_PATH.exists() else None,
    layout="wide",
)

st.markdown(
    """
    <style>
    [data-testid="stDeployButton"] { display: none; }
    [data-testid="stMainMenu"] { display: none; }
    .app-header {
        display: flex;
        align-items: center;
        justify-content: flex-start;
        gap: 0.85rem;
        width: 100%;
        padding: 0 0 0.9rem 0;
        margin: 0 0 1.1rem 0;
        border-bottom: 1px solid rgba(148, 163, 184, 0.35);
    }
    .app-header__logo {
        width: 68px !important;
        min-width: 68px !important;
        height: auto;
        display: block;
    }
    .app-header__title {
        margin: 0;
        color: inherit !important;
        font-size: 1.85rem !important;
        font-weight: 700 !important;
        line-height: 1.2 !important;
        letter-spacing: 0 !important;
    }
    .app-header__subtitle {
        margin: 0.18rem 0 0 0;
        color: rgba(148, 163, 184, 0.95) !important;
        font-size: 0.98rem !important;
        font-weight: 400 !important;
        line-height: 1.35 !important;
        letter-spacing: 0 !important;
    }
    @media (max-width: 768px) {
        .app-header {
            gap: 0.7rem;
            padding-bottom: 0.75rem;
        }
        .app-header__logo {
            width: 52px !important;
            min-width: 52px !important;
        }
        .app-header__title {
            font-size: 1.45rem !important;
        }
        .app-header__subtitle {
            font-size: 0.88rem !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True,
)

header_html = f"""
<div class="app-header">
    {"<img class='app-header__logo' src='" + LOGO_DATA_URI + "' alt='Логотип'>" if LOGO_DATA_URI else ""}
    <div>
        <div class="app-header__title">{SYSTEM_NAME}</div>
        <div class="app-header__subtitle">{SYSTEM_SUBTITLE}</div>
    </div>
</div>
"""
st.markdown(header_html, unsafe_allow_html=True)


def _normalize_candidates(payload):
    if isinstance(payload, list):
        return [item for item in payload if isinstance(item, dict)]
    if isinstance(payload, dict):
        return [payload]
    return []


def _candidate_full_name(candidate):
    full_name = candidate.get("full_name")
    if full_name:
        return full_name
    parts = [
        candidate.get("last_name"),
        candidate.get("first_name"),
        candidate.get("middle_name"),
    ]
    return " ".join([p for p in parts if p])


def _candidate_phones(candidate):
    phones = [
        candidate.get("phone_mobile"),
        candidate.get("phone_2"),
        candidate.get("phone_3"),
    ]
    phones = [str(p) for p in phones if p]
    if not phones:
        return ""
    if len(phones) == 1:
        return phones[0]
    return f"{phones[0]} ({', '.join(phones[1:])})"


def _candidate_emails(candidate):
    emails = [
        candidate.get("email_1"),
        candidate.get("email_2"),
        candidate.get("email_upgo"),
    ]
    emails = [str(e) for e in emails if e]
    if not emails:
        return ""
    if len(emails) == 1:
        return emails[0]
    return f"{emails[0]} ({', '.join(emails[1:])})"


FIELD_LABELS = {
    "date_received": "Дата получения",
    "previous_last_name": "Предыдущая фамилия",
    "sex": "Пол",
    "birth_place": "Место рождения",
    "snils": "СНИЛС",
    "passport_number": "Паспорт",
    "passport_issued": "Кем выдан паспорт",
    "appointment_date": "Дата назначения",
    "dismissal_date": "Дата увольнения",
    "confirmed_experience_years": "Подтвержденный стаж, лет",
    "source_info": "Источник",
    "education_text": "Образование",
    "education_count": "Количество образований",
    "work_text": "Опыт работы",
    "extra_info_text": "Дополнительная информация",
    "status": "Статус",
    "ready_to_work": "Готовность к работе",
    "citizenship": "Гражданство",
}

PRIMARY_FIELDS = {
    "full_name",
    "last_name",
    "first_name",
    "middle_name",
    "id",
    "birth_date",
    "residence_area",
    "phone_mobile",
    "phone_2",
    "phone_3",
    "email_1",
    "email_2",
    "email_upgo",
}


def _display_value(value):
    if value is None or value == "":
        return "—"
    return str(value)


def _candidate_filename(candidate, suffix=".docx"):
    name = _candidate_full_name(candidate) or "candidate"
    name = re.sub(r"[^\wа-яА-ЯёЁ -]+", "", name, flags=re.UNICODE).strip()
    name = re.sub(r"\s+", "_", name)
    return f"{name or 'candidate'}{suffix}"


def _add_doc_paragraph(document, label, value, bold_label=True):
    paragraph = document.add_paragraph()
    if bold_label:
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
        paragraph.add_run(_display_value(value))
    else:
        paragraph.add_run(_display_value(value))
    return paragraph


def _add_doc_table(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = _display_value(value)
        if cells[0].paragraphs and cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].bold = True
    return table


def _set_paragraph_text(paragraph, text):
    text = _display_value(text)
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _set_cell_text(cell, text):
    text = _display_value(text)
    if cell.paragraphs:
        _set_paragraph_text(cell.paragraphs[0], text)
        for paragraph in cell.paragraphs[1:]:
            _set_paragraph_text(paragraph, "")
    else:
        cell.text = text


def _split_education(candidate):
    education = candidate.get("education_text") or ""
    if not education:
        return "—", "—", "—", "—"

    degree = "—"
    if "бакалавр" in education.lower():
        degree = "бакалавр"
    elif "магистр" in education.lower():
        degree = "магистр"
    elif "специалист" in education.lower():
        degree = "специалист"

    direction = "—"
    match = re.search(r'направлен[а-я ]*[:"]+\s*"?([^";]+)', education, flags=re.IGNORECASE)
    if match:
        direction = match.group(1).strip()
    else:
        match = re.search(r"специальност[ьи]:?\s*([^;]+)", education, flags=re.IGNORECASE)
        if match:
            direction = match.group(1).strip()

    return education, "не имеет", education, direction


def _fill_candidate_template(document, candidate):
    full_name = _candidate_full_name(candidate) or "Фамилия Имя Отчество"
    phone = _candidate_phones(candidate)
    email = _candidate_emails(candidate)
    education, academic_degree, graduated, direction = _split_education(candidate)

    if len(document.tables) >= 1:
        header_cell = document.tables[0].rows[0].cells[1]
        if len(header_cell.paragraphs) > 2:
            _set_paragraph_text(header_cell.paragraphs[2], full_name)
        if len(header_cell.paragraphs) > 5:
            _set_paragraph_text(header_cell.paragraphs[5], f"контактный телефон: {_display_value(phone)}")
        if len(header_cell.paragraphs) > 6:
            _set_paragraph_text(header_cell.paragraphs[6], f"адрес эл. почты: {_display_value(email)}")

    if len(document.tables) >= 2 and len(document.tables[1].rows) > 1:
        row = document.tables[1].rows[1]
        _set_cell_text(row.cells[0], candidate.get("birth_date"))
        _set_cell_text(row.cells[1], candidate.get("birth_place"))

    if len(document.tables) >= 3 and len(document.tables[2].rows) > 1:
        row = document.tables[2].rows[1]
        _set_cell_text(row.cells[0], education)
        _set_cell_text(row.cells[1], academic_degree)

    if len(document.tables) >= 4 and len(document.tables[3].rows) > 1:
        row = document.tables[3].rows[1]
        _set_cell_text(row.cells[0], graduated)
        _set_cell_text(row.cells[1], direction)

    if len(document.tables) >= 5 and len(document.tables[4].rows) > 1:
        row = document.tables[4].rows[1]
        _set_cell_text(row.cells[0], "—")
        _set_cell_text(row.cells[1], "—")

    if len(document.tables) >= 6:
        row = document.tables[5].rows[0]
        if len(row.cells) >= 2:
            _set_cell_text(row.cells[0], "Государственные и ведомственные награды, почетные звания\n\n—")
            rank = candidate.get("status") or "—"
            _set_cell_text(row.cells[1], f"Классный чин, звание\n\n{rank}")

    if len(document.tables) >= 7 and len(document.tables[6].rows) > 1:
        row = document.tables[6].rows[1]
        work_text = candidate.get("work_text") or "—"
        if len(row.cells) >= 5:
            _set_cell_text(row.cells[0], "—")
            _set_cell_text(row.cells[1], work_text)
            _set_cell_text(row.cells[3], candidate.get("residence_area"))

    replacements = {
        "Семейное положение:": "Семейное положение: —",
        "Зарегистрирован:": f"Зарегистрирован: {_display_value(candidate.get('residence_area'))}",
        "Фактическое проживание:": f"Фактическое проживание: {_display_value(candidate.get('residence_area'))}",
        "Дополнительная информация:": f"Дополнительная информация: {_display_value(candidate.get('extra_info_text'))}",
        "ИНН:": "ИНН: —",
    }
    for paragraph in document.paragraphs:
        for prefix, value in replacements.items():
            if paragraph.text.strip().startswith(prefix):
                _set_paragraph_text(paragraph, value)
                break


def _build_candidate_docx(candidate):
    if PROFILE_TEMPLATE_PATH.exists():
        document = Document(PROFILE_TEMPLATE_PATH)
        _fill_candidate_template(document, candidate)
    else:
        document = Document()
        styles = document.styles
        styles["Normal"].font.name = "Times New Roman"
        styles["Normal"].font.size = Pt(12)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("СПРАВКА")
        title_run.bold = True
        title_run.font.size = Pt(14)
        _add_doc_paragraph(document, "ФИО", _candidate_full_name(candidate))
        _add_doc_paragraph(document, "контактный телефон", _candidate_phones(candidate))
        _add_doc_paragraph(document, "адрес эл. почты", _candidate_emails(candidate))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _build_candidates_zip(candidates):
    buffer = io.BytesIO()
    used_names = set()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for index, candidate in enumerate(candidates, start=1):
            filename = _candidate_filename(candidate)
            if filename in used_names:
                base = filename[:-5]
                filename = f"{base}_{index}.docx"
            used_names.add(filename)
            archive.writestr(filename, _build_candidate_docx(candidate))
    buffer.seek(0)
    return buffer.getvalue()


def _render_candidates(candidates, key_prefix="candidates"):
    if not candidates:
        return
    st.markdown("**Кандидаты:**")
    for index, candidate in enumerate(candidates):
        with st.container(border=True):
            full_name = _candidate_full_name(candidate)
            st.markdown(f"#### {full_name or '—'}")
            st.markdown(f"**Дата рождения:** {_display_value(candidate.get('birth_date'))}")
            st.markdown(f"**Место проживания:** {_display_value(candidate.get('residence_area'))}")
            st.markdown(f"**Телефон:** {_display_value(_candidate_phones(candidate))}")
            st.markdown(f"**Почта:** {_display_value(_candidate_emails(candidate))}")

            with st.expander("Показать остальные поля"):
                for field, value in candidate.items():
                    if field in PRIMARY_FIELDS:
                        continue
                    st.markdown(f"**{FIELD_LABELS.get(field, field)}:** {_display_value(value)}")

            st.download_button(
                "Скачать анкету",
                data=_build_candidate_docx(candidate),
                file_name=_candidate_filename(candidate),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"{key_prefix}_candidate_docx_{candidate.get('id') or index}",
            )

    st.download_button(
        "Скачать все анкеты",
        data=_build_candidates_zip(candidates),
        file_name="ankety_kandidatov.zip",
        mime="application/zip",
        key=f"{key_prefix}_all_candidates_docx",
    )


def _poll_candidates(session_id):
    if not SERVER_URL or not session_id:
        return [], False
    try:
        base = (SERVER_URL or "").rstrip("/")
        url = f"{base}/session/{session_id}/candidates/current"
        r = requests.get(url, timeout=10)
        r.raise_for_status()
        r_json = r.json()
        return _normalize_candidates(r_json.get("candidates", []) or []), bool(r_json.get("pending"))
    except Exception:
        return [], False


def _post_chat(payload):
    r = requests.post(SERVER_URL, json=payload)
    r.raise_for_status()
    return r.json()


def _post_stop(session_id):
    if not SERVER_URL or not session_id:
        return False
    base = (SERVER_URL or "").rstrip("/")
    url = f"{base}/session/{session_id}/stop"
    r = requests.post(url, timeout=10)
    r.raise_for_status()
    return True


def _finalize_pending_response():
    pending_future = st.session_state.pending_future
    if not pending_future or not pending_future.done():
        return False
    reply = ""
    candidates = []
    try:
        r_json = pending_future.result()
        st.session_state.session_id = r_json.get("session_id", st.session_state.session_id)
        reply = r_json.get("answer", "")
        candidates = _normalize_candidates(r_json.get("candidates", []) or [])
        if not candidates:
            polled_once, _ = _poll_candidates(st.session_state.session_id)
            if polled_once:
                candidates = _normalize_candidates(polled_once)
        if not candidates:
            candidates = _normalize_candidates(st.session_state.last_polled_candidates or [])
    except Exception as e:
        reply = f"Ошибка при запросе к серверу: {e}"
        candidates = []
    st.session_state.pending_future = None
    st.session_state.stop_requested = False
    if reply or candidates:
        final_message = {
            "role": "assistant",
            "content": reply,
            "candidates": candidates,
        }
        st.session_state.messages.append(final_message)
        st.session_state.final_fragment_message = final_message
    else:
        st.session_state.final_fragment_message = None
    return True


# ---------------- LEFT: CHAT (основная область) ----------------
st.subheader("Чат")

_finalize_pending_response()

pending_future = st.session_state.pending_future
pending_in_flight = pending_future is not None and not pending_future.done()

for message_index, m in enumerate(st.session_state.messages):
    content = m["content"]
    if m["role"] == "assistant" and "tool_call" in content:
        continue
    with st.chat_message(m["role"]):
        if m["role"] == "assistant":
            st.markdown(content)
            _render_candidates(m.get("candidates") or [], key_prefix=f"message_{message_index}")
        else:
            st.markdown(content)

if pending_in_flight:
    @st.fragment(run_every=POLL_INTERVAL)
    def _search_status():
        pending_future = st.session_state.pending_future
        if not pending_future:
            final_message = st.session_state.final_fragment_message
            if isinstance(final_message, dict):
                with st.chat_message("assistant"):
                    st.markdown(final_message.get("content") or "")
                    _render_candidates(final_message.get("candidates") or [], key_prefix="final_fragment")
            return
        if pending_future.done():
            _finalize_pending_response()
            final_message = st.session_state.final_fragment_message
            if isinstance(final_message, dict):
                with st.chat_message("assistant"):
                    st.markdown(final_message.get("content") or "")
                    _render_candidates(final_message.get("candidates") or [], key_prefix="final_done")
            st.session_state.last_polled_candidates = []
            return
        polled, is_pending = _poll_candidates(st.session_state.session_id)
        st.session_state.last_polled_candidates = polled
        with st.chat_message("assistant"):
            _render_candidates(polled, key_prefix="polling")
            if not is_pending:
                return
            stop_disabled = st.session_state.stop_requested
            if st.button("Остановить поиск", disabled=stop_disabled, key="stop_search"):
                try:
                    if _post_stop(st.session_state.session_id):
                        st.session_state.stop_requested = True
                except Exception:
                    st.session_state.stop_requested = False
            if st.session_state.stop_requested:
                st.markdown("Останавливаю поиск...")
            else:
                st.markdown("Ищу по базе...")

    _search_status()

prompt = st.chat_input("Введите запрос")

if prompt:
    active_future = st.session_state.pending_future
    if active_future is not None and not active_future.done():
        st.session_state.messages.append(
            {
                "role": "assistant",
                "content": "Поиск еще выполняется. Дождитесь завершения текущего запроса или остановите поиск.",
            }
        )
        st.rerun()

    st.session_state.messages.append({"role": "user", "content": prompt})

    if not SERVER_URL:
        st.session_state.messages.append(
            {
                "role": "assistant",
                "content": "SERVER_URL не задан в переменных окружения.",
            }
        )
    else:
        try:
            if not st.session_state.session_id:
                st.session_state.session_id = str(uuid.uuid4())
            payload = {"session_id": st.session_state.session_id, "message": prompt}
            st.session_state.last_polled_candidates = []
            st.session_state.final_fragment_message = None
            st.session_state.pending_future = EXECUTOR.submit(_post_chat, payload)
            st.session_state.stop_requested = False
        except Exception as e:
            st.session_state.messages.append(
                {
                    "role": "assistant",
                    "content": f"Ошибка при запросе к серверу: {e}",
                }
            )

    st.rerun()
